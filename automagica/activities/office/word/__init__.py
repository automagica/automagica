def open_document(file_path=None):
    """Opens a Word document for further manipulation

    :param file_path: Path to the Word document file (relative or absolute)
    """
    from docx import Document

    return Document(file_path)


def read_text(document):
    """Read the text from a Word document

    :param document: A handle to the document returned by the :function:`open_document` function

    :return: All the text from the document as a string
    """
    text = ""
    for paragraph in document.paragraphs:
        text += paragraph.text + "\r\n"
    return text


def add_heading(document, text="", level=1):
    """Add a heading to a Word document

    :param document: A handle to the document returned by the :function:`open_document` function
    :param text: The text for the heading
    :param level: The level of the heading (defaults to 1)
    """
    document.add_heading(text, level)


def add_paragraph(document, text=""):
    """Add a paragraph to a Word document

    :param document: A handle to the document returned by the :function:`open_document` function
    :param text: The text for the heading
    """
    document.add_paragraph(text)


def add_page_break(document):
    """Add a page break to a Word document
    """
    document.add_page_break()


def add_picture(document, image_path, width=None, height=None):
    """Add a picture to a Word document

    :param document: A handle to the document returned by the :function:`open_document` function
    :param image_path: Path to the image (relative or absolute)
    :param width: Width of the image in pixels
    :param height: Height of the image in pixels
    """
    document.add_picture(image_path, width=width, height=height)


def add_table(document, dataframe, style=None):
    """Add a pandas dataframe as a table to a Word document
    
    :param document: A handle to the document returned by the :function:`open_document` function
    :param dataframe: A :class:`pandas.DataFrame`
    :param style: Name of the table style to apply. Warning: add a table with this style to the document first, then save the document, then remove the table. This way the style is available.
    """
    rows, cols = dataframe.shape

    table = document.add_table(rows=rows, cols=cols, style=style)

    # Table header cells
    header_cells = table.rows[0].cells
    for i, column in enumerate(dataframe.columns):
        header_cells[i].text = str(column)

    # Table rows
    for i, row in dataframe.iterrows():
        row_cells = table.add_row().cells
        for j, column in enumerate(dataframe.columns):
            row_cells[j].text = str(row[column])


def replace_text(document, text=None, replace_with=None):
    """Replace text in a Word document

    :param document: A handle to the document returned by the :function:`open_document` function
    :param text: Text to replace
    :param replace_with: Text to replace with
    """
    for paragraph in document.paragraphs:
        paragraph.text = paragraph.text.replace(text, replace_with)
    return document


def convert_to_pdf(input_path=None, output_path=None):
    """Convert a Word document to a PDF document

    :param input_path: Path to the Word document
    :param output_path: Path to the output PDF document
    """
    import platform
    from win32com import client

    if platform.system() == "Windows":
        word = client.DispatchEx("Word.Application")
        document = word.Documents.Open(input_path)
        document.SaveAs(output_path, FileFormat=17)
        document.Close()
    else:
        raise NotImplementedError(
            "Other OS than Windows are currently not implemented for this function."
        )

